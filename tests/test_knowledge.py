from __future__ import annotations

from collections.abc import Iterator
from contextlib import suppress
from pathlib import Path

import pytest
from fastapi.testclient import TestClient
from sqlmodel import Session, select

from zengrowth.api.main import app
from zengrowth.db import get_session
from zengrowth.knowledge.embeddings import EmbeddingClient
from zengrowth.knowledge.extractor import (
    ExtractedClaim,
    ExtractedEntity,
    ExtractedRelationship,
    ExtractionResult,
    KnowledgeExtractionClient,
)
from zengrowth.knowledge.local_graph import build_local_graph
from zengrowth.knowledge.parsers import parse_document
from zengrowth.knowledge.service import (
    activate_version,
    active_cv_template_text,
    diff_source_versions,
    file_sha256,
    ingest_path,
    paste_document,
    set_active_template,
    summarize_version_diff,
)
from zengrowth.materials.generator import _load_evidence, _read_cv_template
from zengrowth.models import (
    ClaimVerificationState,
    EvidenceClaim,
    SourceDocument,
    SourceDocumentStatus,
    SourceDocumentType,
)


class FakeExtractor(KnowledgeExtractionClient):
    def extract(self, *, text: str, metadata: dict[str, object], model: str) -> ExtractionResult:
        return ExtractionResult(
            claims=[
                ExtractedClaim(
                    claim_text="Led a GraphRAG investment agent project.",
                    category="technical",
                    confidence=0.8,
                    source_span="GraphRAG investment agent",
                    tags=["graphrag", "agent"],
                ),
                ExtractedClaim(
                    claim_text="Unspanned claim.",
                    category="general",
                    confidence=0.95,
                    source_span=None,
                ),
            ],
            entities=[
                ExtractedEntity(name="GraphRAG investment agent", entity_type="project"),
                ExtractedEntity(name="Neo4j", entity_type="tool"),
            ],
            relationships=[
                ExtractedRelationship(
                    source="GraphRAG investment agent",
                    target="Neo4j",
                    relationship_type="USED",
                    confidence=0.9,
                )
            ],
        )


class FakeEmbedder(EmbeddingClient):
    def embed(self, texts: list[str]) -> list[list[float]]:
        return [[0.1, 0.2, 0.3] for _ in texts]


class FailingExtractor(KnowledgeExtractionClient):
    def extract(self, *, text: str, metadata: dict[str, object], model: str) -> ExtractionResult:
        raise ValueError("extractor failed")


def test_parse_text_markdown_latex_and_docx(tmp_path: Path):
    md = tmp_path / "project.md"
    md.write_text("# Project\n\nBuilt a graph.", encoding="utf-8")
    txt = tmp_path / "note.txt"
    txt.write_text("Plain evidence.", encoding="utf-8")
    tex = tmp_path / "cv.tex"
    tex.write_text(r"\section*{Experience}\n\item Led AI delivery", encoding="utf-8")

    from docx import Document

    docx = tmp_path / "cv.docx"
    document = Document()
    document.add_paragraph("DOCX evidence paragraph.")
    document.save(docx)

    assert "Built a graph" in parse_document(md).text
    assert "Plain evidence" in parse_document(txt).text
    assert "Led AI delivery" in parse_document(tex).text
    assert "DOCX evidence" in parse_document(docx).text


def test_parse_pdf_extracts_text(tmp_path: Path):
    pdf = tmp_path / "sample.pdf"
    pdf.write_bytes(
        b"%PDF-1.4\n"
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n"
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n"
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 300 144] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >> endobj\n"
        b"4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n"
        b"5 0 obj << /Length 58 >> stream\n"
        b"BT /F1 18 Tf 50 100 Td (PDF evidence paragraph.) Tj ET\n"
        b"endstream endobj\n"
        b"xref\n0 6\n0000000000 65535 f \n0000000009 00000 n \n"
        b"0000000058 00000 n \n0000000115 00000 n \n0000000241 00000 n \n"
        b"0000000311 00000 n \ntrailer << /Root 1 0 R /Size 6 >>\nstartxref\n419\n%%EOF\n"
    )

    assert "PDF evidence" in parse_document(pdf).text


def test_ingest_path_creates_verified_and_draft_claims(session: Session, tmp_path: Path, monkeypatch):
    root = tmp_path / "knowledge"
    monkeypatch.setenv("KNOWLEDGE_ROOT", str(root))
    source = tmp_path / "project.txt"
    source.write_text("GraphRAG investment agent used Neo4j.", encoding="utf-8")
    result = ingest_path(
        session,
        source,
        source_type=SourceDocumentType.project,
        extractor=FakeExtractor(),
        embedder=FakeEmbedder(),
    )

    assert result.created is True
    assert result.claims == 2
    assert result.verified_claims == 1
    claims = session.exec(select(EvidenceClaim)).all()
    states = {claim.claim_text: claim.verification_state for claim in claims}
    assert states["Led a GraphRAG investment agent project."] == ClaimVerificationState.verified
    assert states["Unspanned claim."] == ClaimVerificationState.draft

    duplicate = ingest_path(
        session,
        source,
        source_type=SourceDocumentType.project,
        extractor=FakeExtractor(),
        embedder=FakeEmbedder(),
    )
    assert duplicate.created is False
    original = session.get(SourceDocument, result.document.id)
    assert original is not None
    assert original.status == SourceDocumentStatus.extracted


def test_embeddings_skipped_by_default(session: Session, tmp_path: Path, monkeypatch):
    """Embeddings are opt-in: with no embedder and the flag off, chunks store no
    vectors and no embedding call is made (RET-01 — the index is unused)."""
    from zengrowth.models import SourceChunk

    root = tmp_path / "knowledge"
    monkeypatch.setenv("KNOWLEDGE_ROOT", str(root))
    source = tmp_path / "project.txt"
    source.write_text("GraphRAG investment agent used Neo4j.", encoding="utf-8")

    result = ingest_path(
        session,
        source,
        source_type=SourceDocumentType.project,
        extractor=FakeExtractor(),
        # no embedder supplied; knowledge_embeddings_enabled defaults to False
    )

    chunks = session.exec(
        select(SourceChunk).where(SourceChunk.source_document_id == result.document.id)
    ).all()
    assert chunks
    assert all(chunk.embedding is None for chunk in chunks)


def test_seed_documents_ingests_portfolio_directory(session: Session, tmp_path: Path, monkeypatch):
    """Portfolio docs in a directory become knowledge sources viewable on the graph."""
    from zengrowth.knowledge.seed import seed_documents
    from zengrowth.models import SourceDocument

    monkeypatch.setenv("KNOWLEDGE_ROOT", str(tmp_path / "knowledge"))
    portfolio = tmp_path / "career"
    portfolio.mkdir()
    (portfolio / "project-impacts.md").write_text(
        "# Project Impacts\n\nBuilt a GraphRAG investment agent using Neo4j.", encoding="utf-8"
    )

    results = seed_documents(session, portfolio, extractor=FakeExtractor())

    assert len(results) == 1
    assert results[0].created is True
    docs = session.exec(select(SourceDocument)).all()
    assert any(doc.filename == "project-impacts.md" for doc in docs)


def test_seed_documents_missing_directory_is_noop(session: Session, tmp_path: Path):
    from zengrowth.knowledge.seed import seed_documents

    assert seed_documents(session, tmp_path / "does-not-exist") == []


class HallucinatedSpanExtractor(KnowledgeExtractionClient):
    """High confidence + a source_span that does NOT appear in the source text."""

    def extract(self, *, text: str, metadata: dict[str, object], model: str) -> ExtractionResult:
        return ExtractionResult(
            claims=[
                ExtractedClaim(
                    claim_text="Increased revenue by 40%.",
                    category="delivery",
                    confidence=0.99,
                    source_span="increased revenue by 40 percent year over year",
                )
            ],
        )


def test_high_confidence_claim_with_unsupported_span_is_not_auto_verified(
    session: Session, tmp_path: Path, monkeypatch
):
    """TP-02: a span the model invented must not auto-verify, even at conf 0.99."""
    root = tmp_path / "knowledge"
    monkeypatch.setenv("KNOWLEDGE_ROOT", str(root))
    source = tmp_path / "project.txt"
    source.write_text("GraphRAG investment agent used Neo4j.", encoding="utf-8")
    result = ingest_path(
        session,
        source,
        source_type=SourceDocumentType.project,
        extractor=HallucinatedSpanExtractor(),
        embedder=FakeEmbedder(),
    )
    assert result.verified_claims == 0
    claim = session.exec(select(EvidenceClaim)).one()
    assert claim.verification_state == ClaimVerificationState.draft


def test_duplicate_failed_source_retries_existing_document(
    session: Session,
    tmp_path: Path,
    monkeypatch,
):
    root = tmp_path / "knowledge"
    monkeypatch.setenv("KNOWLEDGE_ROOT", str(root))
    source = tmp_path / "project.txt"
    source.write_text("GraphRAG investment agent used Neo4j.", encoding="utf-8")
    digest = file_sha256(source)
    original_path = root / "originals" / f"{digest}.txt"
    original_path.parent.mkdir(parents=True)
    original_path.write_text(source.read_text(encoding="utf-8"), encoding="utf-8")
    failed = SourceDocument(
        filename=source.name,
        original_path=str(original_path),
        content_hash=digest,
        source_type=SourceDocumentType.project,
        status=SourceDocumentStatus.failed,
        error="OPENAI_API_KEY is required",
    )
    session.add(failed)
    session.commit()
    session.refresh(failed)

    result = ingest_path(
        session,
        source,
        source_type=SourceDocumentType.project,
        extractor=FakeExtractor(),
        embedder=FakeEmbedder(),
    )

    assert result.created is False
    assert result.document.id == failed.id
    assert result.document.status == SourceDocumentStatus.extracted
    assert result.document.error is None
    assert result.claims == 2


def test_duplicate_failed_source_records_retry_failure(
    session: Session,
    tmp_path: Path,
    monkeypatch,
):
    root = tmp_path / "knowledge"
    monkeypatch.setenv("KNOWLEDGE_ROOT", str(root))
    source = tmp_path / "project.txt"
    source.write_text("GraphRAG investment agent used Neo4j.", encoding="utf-8")
    digest = file_sha256(source)
    original_path = root / "originals" / f"{digest}.txt"
    original_path.parent.mkdir(parents=True)
    original_path.write_text(source.read_text(encoding="utf-8"), encoding="utf-8")
    failed = SourceDocument(
        filename=source.name,
        original_path=str(original_path),
        content_hash=digest,
        source_type=SourceDocumentType.project,
        status=SourceDocumentStatus.failed,
        error="old error",
    )
    session.add(failed)
    session.commit()
    session.refresh(failed)

    with suppress(ValueError):
        ingest_path(
            session,
            source,
            source_type=SourceDocumentType.project,
            extractor=FailingExtractor(),
            embedder=FakeEmbedder(),
        )

    session.refresh(failed)
    assert failed.status == SourceDocumentStatus.failed
    assert failed.error == "extractor failed"


class _FailOnSecondChunkExtractor(KnowledgeExtractionClient):
    """Succeeds on the first chunk (committing claims), fails on the second."""

    def __init__(self) -> None:
        self.calls = 0

    def extract(self, *, text: str, metadata: dict[str, object], model: str) -> ExtractionResult:
        self.calls += 1
        if self.calls >= 2:
            raise ValueError("extractor failed on later chunk")
        return ExtractionResult(
            claims=[
                ExtractedClaim(
                    claim_text="Partial claim from the first chunk.",
                    category="technical",
                    confidence=0.9,
                    source_span="first chunk",
                )
            ],
            entities=[ExtractedEntity(name="Partial Entity", entity_type="project")],
            relationships=[],
        )


def test_failed_ingest_leaves_no_partial_derivatives(
    session: Session, tmp_path: Path, monkeypatch
):
    """EA-02: a mid-document failure must not leave orphan chunks/claims/entities."""
    from zengrowth.models import KnowledgeEntity, SourceChunk

    root = tmp_path / "knowledge"
    monkeypatch.setenv("KNOWLEDGE_ROOT", str(root))
    source = tmp_path / "long.txt"
    # Two ~1.5k-char paragraphs → two chunks; extraction commits the first, then fails.
    para = "GraphRAG investment agent used Neo4j to ground claims. " * 30
    source.write_text(f"{para}\n\n{para}", encoding="utf-8")

    extractor = _FailOnSecondChunkExtractor()
    with suppress(ValueError):
        ingest_path(
            session,
            source,
            source_type=SourceDocumentType.project,
            extractor=extractor,
            embedder=FakeEmbedder(),
        )
    assert extractor.calls >= 2  # the failure happened after the first chunk committed

    doc = session.exec(select(SourceDocument)).first()
    assert doc is not None
    assert doc.status == SourceDocumentStatus.failed
    # No orphan derivatives survive the failed document.
    assert session.exec(select(SourceChunk).where(SourceChunk.source_document_id == doc.id)).all() == []
    assert session.exec(select(EvidenceClaim).where(EvidenceClaim.source_document_id == doc.id)).all() == []
    assert session.exec(select(KnowledgeEntity).where(KnowledgeEntity.source_document_id == doc.id)).all() == []


def test_verified_claims_feed_material_evidence(session: Session):
    claim = EvidenceClaim(
        id="claim-test",
        source_document_id=1,
        claim_text="Built a production GraphRAG agent.",
        category="technical",
        confidence=0.9,
        verification_state=ClaimVerificationState.verified,
        source_span="source span",
        tags=["graphrag"],
    )
    session.add(claim)
    session.commit()

    evidence = _load_evidence(session)

    assert evidence[0].id == "claim-test"
    assert evidence[0].claim_text == "Built a production GraphRAG agent."


def _fake_ingest_kwargs() -> dict[str, object]:
    return {
        "extractor": FakeExtractor(),
        "embedder": FakeEmbedder(),
    }


def test_paste_document_creates_versioned_source(session: Session, tmp_path: Path, monkeypatch):
    monkeypatch.setenv("KNOWLEDGE_ROOT", str(tmp_path / "knowledge"))

    first = paste_document(
        session,
        text="GraphRAG investment agent used Neo4j. Version one.",
        filename="cv-style",
        fmt="tex",
        source_type=SourceDocumentType.cv,
        title="LaTeX CV style",
        **_fake_ingest_kwargs(),
    )
    assert first.created is True
    assert first.document.version == 1
    assert first.document.is_current is True
    assert first.document.title == "LaTeX CV style"
    assert first.document.summary  # derived, non-empty

    second = paste_document(
        session,
        text="GraphRAG investment agent used Neo4j. Version two, expanded.",
        filename="cv-style",
        fmt="tex",
        source_type=SourceDocumentType.cv,
        supersedes_id=first.document.id,
        **_fake_ingest_kwargs(),
    )
    assert second.document.version == 2
    assert second.document.supersedes_id == first.document.id
    assert second.document.lineage_id == first.document.lineage_id
    assert second.document.is_current is True

    prior = session.get(SourceDocument, first.document.id)
    assert prior is not None
    assert prior.is_current is False


def test_paste_rejects_unknown_format_and_empty(session: Session, tmp_path: Path, monkeypatch):
    monkeypatch.setenv("KNOWLEDGE_ROOT", str(tmp_path / "knowledge"))

    with pytest.raises(ValueError):
        paste_document(session, text="x", filename="f", fmt="pdf", **_fake_ingest_kwargs())
    with pytest.raises(ValueError):
        paste_document(session, text="   ", filename="f", fmt="tex", **_fake_ingest_kwargs())


def test_template_promotion_drives_cv_generator(session: Session, tmp_path: Path, monkeypatch):
    monkeypatch.setenv("KNOWLEDGE_ROOT", str(tmp_path / "knowledge"))
    tex = r"\section*{Professional Summary} Active promoted template body."

    # No promotion yet -> generator falls back to checked-in cv_source.tex.
    assert active_cv_template_text(session) is None

    result = paste_document(
        session,
        text=tex,
        filename="cv-style",
        fmt="tex",
        source_type=SourceDocumentType.cv,
        promote_template=True,
        **_fake_ingest_kwargs(),
    )
    assert result.document.template_role == "cv_style"
    assert result.document.is_current is True

    active = active_cv_template_text(session)
    assert active is not None
    assert "Active promoted template body." in active
    assert "Active promoted template body." in _read_cv_template(session)


def test_tex_upload_auto_promotes_and_supersedes_prior_style(
    session: Session, tmp_path: Path, monkeypatch
):
    monkeypatch.setenv("KNOWLEDGE_ROOT", str(tmp_path / "knowledge"))

    first_path = tmp_path / "cv-style-v1.tex"
    first_path.write_text(r"\section*{Professional Summary} First style body.", encoding="utf-8")
    first = ingest_path(
        session,
        first_path,
        source_type=SourceDocumentType.cv,
        promote_template=True,
        **_fake_ingest_kwargs(),
    )
    assert first.document.template_role == "cv_style"
    assert "First style body." in (active_cv_template_text(session) or "")

    # A newly uploaded .tex style supersedes the prior one and becomes active.
    second_path = tmp_path / "cv-style-v2.tex"
    second_path.write_text(r"\section*{Professional Summary} Second style body.", encoding="utf-8")
    second = ingest_path(
        session,
        second_path,
        source_type=SourceDocumentType.cv,
        promote_template=True,
        **_fake_ingest_kwargs(),
    )
    assert second.document.template_role == "cv_style"
    assert second.document.is_current is True
    assert second.document.supersedes_id == first.document.id

    prior = session.get(SourceDocument, first.document.id)
    assert prior is not None and prior.is_current is False
    assert "Second style body." in (active_cv_template_text(session) or "")
    assert "Second style body." in _read_cv_template(session)


def test_activate_version_rolls_back_current_head(session: Session, tmp_path: Path, monkeypatch):
    monkeypatch.setenv("KNOWLEDGE_ROOT", str(tmp_path / "knowledge"))
    first = paste_document(
        session, text="One.", filename="doc", fmt="md", **_fake_ingest_kwargs()
    )
    second = paste_document(
        session,
        text="Two.",
        filename="doc",
        fmt="md",
        supersedes_id=first.document.id,
        **_fake_ingest_kwargs(),
    )
    assert second.document.is_current is True

    activate_version(session, first.document.id or 0)
    rolled_first = session.get(SourceDocument, first.document.id)
    rolled_second = session.get(SourceDocument, second.document.id)
    assert rolled_first is not None and rolled_first.is_current is True
    assert rolled_second is not None and rolled_second.is_current is False


def test_diff_source_versions_reports_added_and_removed(session: Session, tmp_path: Path, monkeypatch):
    monkeypatch.setenv("KNOWLEDGE_ROOT", str(tmp_path / "knowledge"))
    first = paste_document(
        session, text="alpha\nbeta\ngamma", filename="doc", fmt="md", **_fake_ingest_kwargs()
    )
    second = paste_document(
        session,
        text="alpha\nbeta updated\ngamma\ndelta",
        filename="doc",
        fmt="md",
        supersedes_id=first.document.id,
        **_fake_ingest_kwargs(),
    )

    result = diff_source_versions(session, first.document.id or 0, second.document.id or 0)
    assert result["base_version"] == 1
    assert result["target_version"] == 2
    assert result["added"] >= 2  # "beta updated" + "delta"
    assert result["removed"] >= 1  # "beta"
    ops = {line["op"] for line in result["lines"]}
    assert "add" in ops and "remove" in ops
    added_text = [line["text"] for line in result["lines"] if line["op"] == "add"]
    assert "delta" in added_text


def test_diff_source_versions_rejects_cross_lineage(session: Session, tmp_path: Path, monkeypatch):
    monkeypatch.setenv("KNOWLEDGE_ROOT", str(tmp_path / "knowledge"))
    a = paste_document(session, text="one", filename="a", fmt="md", **_fake_ingest_kwargs())
    b = paste_document(session, text="two", filename="b", fmt="md", **_fake_ingest_kwargs())
    with pytest.raises(ValueError, match="same lineage"):
        diff_source_versions(session, a.document.id or 0, b.document.id or 0)


def test_summarize_version_diff_uses_injected_client(session: Session, tmp_path: Path, monkeypatch):
    monkeypatch.setenv("KNOWLEDGE_ROOT", str(tmp_path / "knowledge"))
    first = paste_document(
        session, text="alpha\nbeta", filename="doc", fmt="md", **_fake_ingest_kwargs()
    )
    second = paste_document(
        session,
        text="alpha\nbeta\ngamma",
        filename="doc",
        fmt="md",
        supersedes_id=first.document.id,
        **_fake_ingest_kwargs(),
    )

    class FakeDiffClient:
        def __init__(self) -> None:
            self.prompts: list[str] = []

        def generate(self, system: str, user: str, model: str, **kwargs) -> dict[str, str]:
            self.prompts.append(user)
            return {"summary": "Added a gamma line."}

    client = FakeDiffClient()
    summary = summarize_version_diff(
        session, first.document.id or 0, second.document.id or 0, client=client
    )
    assert summary == "Added a gamma line."
    assert client.prompts and "gamma" in client.prompts[0]


def test_version_diff_endpoint(monkeypatch, session: Session, tmp_path: Path):
    monkeypatch.setenv("KNOWLEDGE_ROOT", str(tmp_path / "knowledge"))

    def override_get_session() -> Iterator[Session]:
        yield session

    first = paste_document(
        session, text="alpha\nbeta", filename="doc", fmt="md", **_fake_ingest_kwargs()
    )
    second = paste_document(
        session,
        text="alpha\nbeta\ngamma",
        filename="doc",
        fmt="md",
        supersedes_id=first.document.id,
        **_fake_ingest_kwargs(),
    )

    app.dependency_overrides[get_session] = override_get_session
    client = TestClient(app)
    try:
        resp = client.get(
            f"/api/knowledge/sources/{second.document.id}/diff",
            params={"against": first.document.id},
        )
        assert resp.status_code == 200
        body = resp.json()
        assert body["base_version"] == 1 and body["target_version"] == 2
        assert body["added"] >= 1
        assert any(line["op"] == "add" and line["text"] == "gamma" for line in body["lines"])
    finally:
        app.dependency_overrides.clear()


def test_set_active_template_clears_other_lineages(session: Session, tmp_path: Path, monkeypatch):
    monkeypatch.setenv("KNOWLEDGE_ROOT", str(tmp_path / "knowledge"))
    a = paste_document(
        session, text="Template A.", filename="a", fmt="tex", promote_template=True, **_fake_ingest_kwargs()
    )
    b = paste_document(
        session, text="Template B.", filename="b", fmt="tex", **_fake_ingest_kwargs()
    )
    set_active_template(session, b.document.id or 0, role="cv_style")

    refreshed_a = session.get(SourceDocument, a.document.id)
    refreshed_b = session.get(SourceDocument, b.document.id)
    assert refreshed_a is not None and refreshed_a.template_role is None
    assert refreshed_b is not None and refreshed_b.template_role == "cv_style"
    assert active_cv_template_text(session) is not None


def test_build_local_graph_emits_version_and_topic_edges(session: Session, tmp_path: Path, monkeypatch):
    monkeypatch.setenv("KNOWLEDGE_ROOT", str(tmp_path / "knowledge"))
    first = paste_document(
        session, text="GraphRAG investment agent used Neo4j. One.", filename="p1", fmt="md",
        source_type=SourceDocumentType.project, **_fake_ingest_kwargs(),
    )
    paste_document(
        session, text="GraphRAG investment agent used Neo4j. Two.", filename="p1", fmt="md",
        source_type=SourceDocumentType.project, supersedes_id=first.document.id,
        **_fake_ingest_kwargs(),
    )
    # A different lineage that shares entities -> related_to edge.
    paste_document(
        session, text="GraphRAG investment agent used Neo4j. Separate doc.", filename="p2", fmt="md",
        source_type=SourceDocumentType.note, **_fake_ingest_kwargs(),
    )

    graph = build_local_graph(session)
    kinds = {edge.kind for edge in graph.edges}
    assert "supersedes" in kinds
    assert "related_to" in kinds
    assert all(node.kind == "source" for node in graph.nodes)

    expanded = build_local_graph(session, include_claims=True, include_entities=True)
    assert any(node.kind == "claim" for node in expanded.nodes)
    assert any(node.kind == "entity" for node in expanded.nodes)


def test_paste_and_file_endpoints(monkeypatch, session: Session, tmp_path: Path):
    monkeypatch.setenv("KNOWLEDGE_ROOT", str(tmp_path / "knowledge"))

    def override_get_session() -> Iterator[Session]:
        yield session

    # Seed a source on disk we can serve back.
    seeded = paste_document(
        session, text="Servable original body.", filename="served", fmt="md",
        **_fake_ingest_kwargs(),
    )

    app.dependency_overrides[get_session] = override_get_session
    client = TestClient(app)
    try:
        graph = client.get("/api/knowledge/graph")
        assert graph.status_code == 200
        assert "nodes" in graph.json() and "edges" in graph.json()

        served = client.get(f"/api/knowledge/sources/{seeded.document.id}/file", params={"kind": "processed"})
        assert served.status_code == 200
        assert "Servable original body." in served.text
        # Previewable content must render inline, not force a download.
        assert served.headers["content-disposition"].startswith("inline")
        assert served.headers["content-type"].startswith("text/plain")

        original = client.get(f"/api/knowledge/sources/{seeded.document.id}/file", params={"kind": "original"})
        assert original.status_code == 200
        assert original.headers["content-disposition"].startswith("inline")
        assert original.headers["content-type"].startswith("text/plain")

        detail = client.get(f"/api/knowledge/sources/{seeded.document.id}")
        assert detail.status_code == 200
        body = detail.json()
        assert body["version"] == 1
        assert "versions" in body
    finally:
        app.dependency_overrides.clear()


def test_knowledge_api_claim_review(monkeypatch, session: Session, tmp_path: Path):
    def override_get_session() -> Iterator[Session]:
        yield session

    source = SourceDocument(
        filename="project.txt",
        original_path=str(tmp_path / "project.txt"),
        content_hash="hash",
        source_type=SourceDocumentType.project,
    )
    session.add(source)
    session.commit()
    session.refresh(source)
    claim = EvidenceClaim(
        id="claim-api",
        source_document_id=source.id or 0,
        claim_text="Draft claim.",
        category="general",
        confidence=0.6,
        verification_state=ClaimVerificationState.draft,
    )
    session.add(claim)
    session.commit()

    app.dependency_overrides[get_session] = override_get_session
    client = TestClient(app)
    try:
        listed = client.get("/api/knowledge/claims", params={"state": "draft"})
        assert listed.status_code == 200
        assert listed.json()[0]["id"] == "claim-api"

        verified = client.post("/api/knowledge/claims/claim-api/verify")
        assert verified.status_code == 200
        assert verified.json()["verification_state"] == "verified"

        edited = client.patch("/api/knowledge/claims/claim-api", json={"category": "technical"})
        assert edited.status_code == 200
        assert edited.json()["category"] == "technical"
        # Non-substantive edit (category) leaves the prior decision intact.
        assert edited.json()["verification_state"] == "verified"

        # TP-13: editing the claim text of a verified claim reopens it to draft
        # so the new wording is re-reviewed instead of inheriting trust.
        reworded = client.patch(
            "/api/knowledge/claims/claim-api", json={"claim_text": "Reworded claim."}
        )
        assert reworded.status_code == 200
        assert reworded.json()["verification_state"] == "draft"
    finally:
        app.dependency_overrides.clear()


class DistortedClaimExtractor(KnowledgeExtractionClient):
    """High confidence + a REAL span — but the claim inflates the span's figure."""

    def extract(self, *, text: str, metadata: dict[str, object], model: str) -> ExtractionResult:
        return ExtractionResult(
            claims=[
                ExtractedClaim(
                    claim_text="Increased revenue by 40%.",
                    category="delivery",
                    confidence=0.99,
                    source_span="increased revenue by 30%",
                ),
                ExtractedClaim(
                    claim_text="Increased margin by 12%.",
                    category="delivery",
                    confidence=0.99,
                    source_span="increased margin by 12%",
                ),
            ],
        )


def test_distorted_claim_with_real_span_is_not_auto_verified(
    session: Session, tmp_path: Path, monkeypatch
):
    """TP-02b: a genuine span cannot auto-verify a claim whose figures differ from it.

    TP-02 only proved the span exists in the source; the first claim cites a real
    span saying 30% while asserting 40%, so it must stay draft (and be flagged in
    the audit log). The consistent control claim still auto-verifies.
    """
    from zengrowth.models import AuditLog

    root = tmp_path / "knowledge"
    monkeypatch.setenv("KNOWLEDGE_ROOT", str(root))
    source = tmp_path / "project.txt"
    source.write_text(
        "The pricing programme increased revenue by 30% and increased margin by 12%.",
        encoding="utf-8",
    )
    result = ingest_path(
        session,
        source,
        source_type=SourceDocumentType.project,
        extractor=DistortedClaimExtractor(),
        embedder=FakeEmbedder(),
    )
    assert result.claims == 2
    assert result.verified_claims == 1
    states = {
        claim.claim_text: claim.verification_state
        for claim in session.exec(select(EvidenceClaim)).all()
    }
    assert states["Increased revenue by 40%."] == ClaimVerificationState.draft
    assert states["Increased margin by 12%."] == ClaimVerificationState.verified
    flagged = [
        row
        for row in session.exec(select(AuditLog)).all()
        if row.action == "knowledge_claim_distortion_flagged"
    ]
    assert len(flagged) == 1
    assert flagged[0].detail["distortions"] == ["40"]
