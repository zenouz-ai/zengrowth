"""Document parsing for the local knowledge inbox."""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx", ".tex"}


@dataclass
class ParsedDocument:
    text: str
    metadata: dict[str, object]


def parse_document(path: str | Path) -> ParsedDocument:
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"unsupported knowledge document type: {suffix or '<none>'}")
    if suffix in {".md", ".txt"}:
        text = file_path.read_text(encoding="utf-8")
    elif suffix == ".tex":
        text = _parse_tex(file_path.read_text(encoding="utf-8"))
    elif suffix == ".pdf":
        text = _parse_pdf(file_path)
    else:
        text = _parse_docx(file_path)
    text = _normalize_text(text)
    if not text:
        raise ValueError("knowledge document has no extractable text")
    return ParsedDocument(
        text=text,
        metadata={"extension": suffix, "bytes": file_path.stat().st_size},
    )


def _parse_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page_num, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            parts.append(f"[page {page_num}]\n{page_text}")
    return "\n\n".join(parts)


def _parse_docx(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                parts.append(" | ".join(values))
    return "\n".join(parts)


def _parse_tex(text: str) -> str:
    text = re.sub(r"%.*", "", text)
    text = re.sub(r"\\(?:section|subsection|subsubsection)\*?\{([^}]*)\}", r"\n\1\n", text)
    text = re.sub(r"\\(?:textbf|emph|textit)\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\item\s*", "- ", text)
    text = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?(?:\{[^}]*\})?", " ", text)
    text = text.replace(r"\&", "&").replace(r"\%", "%").replace(r"\_", "_")
    return text


def _normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    compact: list[str] = []
    blank = False
    for line in lines:
        if not line:
            if not blank:
                compact.append("")
            blank = True
            continue
        compact.append(line)
        blank = False
    return "\n".join(compact).strip()
